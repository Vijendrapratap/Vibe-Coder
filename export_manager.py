"""
VibeDoc Multi-format Export Manager
Supports Ma# PDF Export
try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# Advanced PDF export - removed weasyprint dependency, using reportlab
WEASYPRINT_AVAILABLE = FalseF Format document export
"""

import os
import io
import re
import zipfile
import tempfile
from datetime import datetime
from typing import Dict, Tuple, Optional, Any
import logging

# Core dependencies
import markdown
import html2text

# Word export
try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# PDF export
try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# Advanced PDF export (backup solution) - removed weasyprint dependency
WEASYPRINT_AVAILABLE = False

logger = logging.getLogger(__name__)

class ExportManager:
    """Multi-format Export Manager"""
    
    def __init__(self):
        self.supported_formats = ['markdown', 'html']
        
        # Check optional dependencies
        if DOCX_AVAILABLE:
            self.supported_formats.append('docx')
        if PDF_AVAILABLE:
            self.supported_formats.append('pdf')
            
        logger.info(f"📄 ExportManager initialized, supported formats: {', '.join(self.supported_formats)}")
    
    def get_supported_formats(self) -> list:
        """Get supported export formats"""
        return self.supported_formats.copy()
    
    def export_to_markdown(self, content: str, metadata: Optional[Dict] = None) -> str:
        """
        Export to Markdown format (clean and optimize)
        
        Args:
            content: Original content
            metadata: Metadata information
            
        Returns:
            str: Optimized Markdown content
        """
        try:
            # Add document header information
            if metadata:
                header = f"""---
title: {metadata.get('title', 'VibeDoc Development Plan')}
author: {metadata.get('author', 'VibeDoc AI Agent')}
date: {metadata.get('date', datetime.now().strftime('%Y-%m-%d'))}
generator: VibeDoc AI Agent v1.0
---

"""
                content = header + content
            
            # Clean and optimize content
            content = self._clean_markdown_content(content)
            
            logger.info("✅ Markdown export successful")
            return content
            
        except Exception as e:
            logger.error(f"❌ Markdown export failed: {e}")
            return content  # Return original content
    
    def export_to_html(self, content: str, metadata: Optional[Dict] = None) -> str:
        """
        Export to HTML format (with styles)
        
        Args:
            content: Markdown content
            metadata: Metadata information
            
        Returns:
            str: Complete HTML content
        """
        try:
            # Configure Markdown extensions
            md = markdown.Markdown(
                extensions=[
                    'markdown.extensions.extra',
                    'markdown.extensions.codehilite',
                    'markdown.extensions.toc',
                    'markdown.extensions.tables'
                ],
                extension_configs={
                    'codehilite': {
                        'css_class': 'highlight',
                        'use_pygments': False
                    },
                    'toc': {
                        'title': 'Table of Contents'
                    }
                }
            )
            
            # Convert Markdown to HTML
            html_content = md.convert(content)
            
            # Generate complete HTML document
            title = metadata.get('title', 'VibeDoc Development Plan') if metadata else 'VibeDoc Development Plan'
            author = metadata.get('author', 'VibeDoc AI Agent') if metadata else 'VibeDoc AI Agent'
            date = metadata.get('date', datetime.now().strftime('%Y-%m-%d')) if metadata else datetime.now().strftime('%Y-%m-%d')
            
            full_html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <meta name="author" content="{author}">
    <meta name="generator" content="VibeDoc AI Agent">
    <style>
        {self._get_html_styles()}
    </style>
    <!-- Mermaid support -->
    <script src="https://cdn.jsdelivr.net/npm/mermaid@10.6.1/dist/mermaid.min.js"></script>
    <script>
        document.addEventListener('DOMContentLoaded', function() {{
            mermaid.initialize({{ 
                startOnLoad: true,
                theme: 'default',
                securityLevel: 'loose',
                flowchart: {{ useMaxWidth: true }}
            }});
        }});
    </script>
</head>
<body>
    <div class="container">
        <header class="document-header">
            <h1>{title}</h1>
            <div class="meta-info">
                <span class="author">📝 {author}</span>
                <span class="date">📅 {date}</span>
                <span class="generator">🤖 Generated by VibeDoc AI Agent</span>
            </div>
        </header>
        
        <main class="content">
            {html_content}
        </main>
        
        <footer class="document-footer">
            <p>This document is generated by <strong>VibeDoc AI Agent</strong> | Generated at: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
        </footer>
    </div>
</body>
</html>"""
            
            logger.info("✅ HTML export successful")
            return full_html
            
        except Exception as e:
            logger.error(f"❌ HTML export failed: {e}")
            # Simple HTML fallback
            return f"""<!DOCTYPE html>
<html><head><title>VibeDoc Development Plan</title></head>
<body><pre>{content}</pre></body></html>"""
    
    def export_to_docx(self, content: str, metadata: Optional[Dict] = None) -> bytes:
        """
        Export as Word document format
        
        Args:
            content: Markdown content
            metadata: Metadata information
            
        Returns:
            bytes: Binary data of the Word document
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is not installed, unable to export Word format")
        
        try:
            # Create new document
            doc = Document()
            
            # Set document properties
            properties = doc.core_properties
            properties.title = metadata.get('title', 'VibeDoc Development Plan') if metadata else 'VibeDoc Development Plan'
            properties.author = metadata.get('author', 'VibeDoc AI Agent') if metadata else 'VibeDoc AI Agent'
            properties.subject = 'AI-driven Intelligent Development Plan'
            properties.comments = 'Generated by VibeDoc AI Agent'
            
            # Add title
            title = doc.add_heading(properties.title, 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add meta information
            doc.add_paragraph()
            meta_para = doc.add_paragraph()
            meta_para.add_run(f"📝 Author: {properties.author}").bold = True
            meta_para.add_run("\n")
            meta_para.add_run(f"📅 Generated at: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}").bold = True
            meta_para.add_run("\n")
            meta_para.add_run("🤖 Generated by: VibeDoc AI Agent").bold = True
            
            doc.add_paragraph()
            doc.add_paragraph("─" * 50)
            doc.add_paragraph()
            
            # Parse and add content
            self._parse_markdown_to_docx(doc, content)
            
            # Add footer
            doc.add_paragraph()
            doc.add_paragraph("─" * 50)
            footer_para = doc.add_paragraph()
            footer_para.add_run("This document is automatically generated by VibeDoc AI Agent").italic = True
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Save to memory
            doc_stream = io.BytesIO()
            doc.save(doc_stream)
            doc_stream.seek(0)
            
            logger.info("✅ Word document export successful")
            return doc_stream.getvalue()
            
        except Exception as e:
            logger.error(f"❌ Word export failed: {e}")
            raise
    
    def export_to_pdf(self, content: str, metadata: Optional[Dict] = None) -> bytes:
        """
        Export as PDF format
        
        Args:
            content: Markdown content  
            metadata: Metadata information
            
        Returns:
            bytes: Binary data of the PDF document
        """
        if PDF_AVAILABLE:
            return self._export_pdf_reportlab(content, metadata)
        else:
            raise ImportError("PDF export dependencies are not installed")
    
    def create_multi_format_export(self, content: str, formats: list = None, metadata: Optional[Dict] = None) -> bytes:
        """
        Create a ZIP package for multi-format export
        
        Args:
            content: Original content
            formats: List of formats to export, defaults to all supported formats
            metadata: Metadata information
            
        Returns:
            bytes: Binary data of the ZIP file
        """
        if formats is None:
            formats = self.supported_formats
        
        # Validate formats
        invalid_formats = set(formats) - set(self.supported_formats)
        if invalid_formats:
            raise ValueError(f"Unsupported formats: {', '.join(invalid_formats)}")
        
        try:
            # Create an in-memory ZIP file
            zip_buffer = io.BytesIO()
            
            with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                # Generate base filename
                base_name = metadata.get('title', 'vibedoc_plan') if metadata else 'vibedoc_plan'
                base_name = re.sub(r'[^\w\-_\.]', '_', base_name)  # Clean filename
                
                # Export various formats
                for fmt in formats:
                    try:
                        if fmt == 'markdown':
                            file_content = self.export_to_markdown(content, metadata)
                            zip_file.writestr(f"{base_name}.md", file_content.encode('utf-8'))
                            
                        elif fmt == 'html':
                            file_content = self.export_to_html(content, metadata)
                            zip_file.writestr(f"{base_name}.html", file_content.encode('utf-8'))
                            
                        elif fmt == 'docx' and DOCX_AVAILABLE:
                            file_content = self.export_to_docx(content, metadata)
                            zip_file.writestr(f"{base_name}.docx", file_content)
                            
                        elif fmt == 'pdf' and PDF_AVAILABLE:
                            file_content = self.export_to_pdf(content, metadata)
                            zip_file.writestr(f"{base_name}.pdf", file_content)
                            
                    except Exception as e:
                        logger.warning(f"⚠️ Export failed for format {fmt}: {e}")
                        # Add error message file to ZIP
                        error_msg = f"Export failed for format {fmt}:\n{str(e)}\n\nPlease check if related dependencies are installed correctly."
                        zip_file.writestr(f"ERROR_{fmt}.txt", error_msg.encode('utf-8'))
                
                # Add README file
                readme_content = f"""# VibeDoc Export Package

## 📋 File Description
This archive contains multiple export formats of your development plan:

### 📄 Supported Formats:
- **Markdown (.md)**: Original format, supports all Markdown syntax
- **HTML (.html)**: Web format, includes styles and Mermaid chart support
- **Word (.docx)**: Microsoft Word document format
- **PDF (.pdf)**: Portable Document Format

### 🤖 Generation Info:
- Generated at: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
- Generated by: VibeDoc AI Agent v1.0
- Project URL: https://github.com/JasonRobertDestiny/VibeDocs

### 💡 Usage Recommendations:
1. Prefer using the HTML format for the best visual experience
2. Use the Markdown format for further editing
3. Use the Word format for formal document processing
4. Use the PDF format for sharing and printing

---
Thank you for using VibeDoc AI Agent!
"""
                zip_file.writestr("README.md", readme_content.encode('utf-8'))
            
            zip_buffer.seek(0)
            logger.info(f"✅ Multi-format export successful, including {len(formats)} formats")
            return zip_buffer.getvalue()
            
        except Exception as e:
            logger.error(f"❌ Multi-format export failed: {e}")
            raise
    
    def _clean_markdown_content(self, content: str) -> str:
        """Clean and optimize Markdown content"""
        # Fix common formatting issues
        content = re.sub(r'\n{3,}', '\n\n', content)  # Remove excessive blank lines
        content = re.sub(r'(?m)^[ \t]+$', '', content)  # Remove lines with only spaces
        content = content.strip()
        
        return content
    
    def _get_html_styles(self) -> str:
        """Get HTML styles"""
        return """
        * {
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }
        
        body {
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', 'PingFang SC', 'Hiragino Sans GB', sans-serif;
            line-height: 1.6;
            color: #333;
            background: #f8fafc;
        }
        
        .container {
            max-width: 900px;
            margin: 0 auto;
            padding: 20px;
            background: white;
            box-shadow: 0 0 20px rgba(0,0,0,0.1);
            border-radius: 8px;
            margin-top: 20px;
            margin-bottom: 20px;
        }
        
        .document-header {
            text-align: center;
            border-bottom: 3px solid #667eea;
            padding-bottom: 20px;
            margin-bottom: 30px;
        }
        
        .document-header h1 {
            color: #667eea;
            font-size: 2.2em;
            margin-bottom: 15px;
        }
        
        .meta-info {
            display: flex;
            justify-content: center;
            gap: 20px;
            flex-wrap: wrap;
            color: #666;
            font-size: 0.9em;
        }
        
        .content h1, .content h2, .content h3, .content h4 {
            color: #2d3748;
            margin-top: 2em;
            margin-bottom: 1em;
        }
        
        .content h1 { border-bottom: 2px solid #667eea; padding-bottom: 0.5em; }
        .content h2 { border-bottom: 1px solid #e2e8f0; padding-bottom: 0.3em; }
        
        .content p {
            margin-bottom: 1em;
            text-align: justify;
        }
        
        .content ul, .content ol {
            margin-bottom: 1em;
            padding-left: 2em;
        }
        
        .content li {
            margin-bottom: 0.5em;
        }
        
        .content pre {
            background: #2d3748;
            color: #e2e8f0;
            padding: 1em;
            border-radius: 6px;
            overflow-x: auto;
            margin: 1em 0;
        }
        
        .content code {
            background: #f7fafc;
            padding: 0.2em 0.4em;
            border-radius: 3px;
            font-family: 'SFMono-Regular', Consolas, monospace;
        }
        
        .content table {
            width: 100%;
            border-collapse: collapse;
            margin: 1em 0;
        }
        
        .content th, .content td {
            border: 1px solid #e2e8f0;
            padding: 0.75em;
            text-align: left;
        }
        
        .content th {
            background: #f7fafc;
            font-weight: 600;
        }
        
        .content blockquote {
            border-left: 4px solid #667eea;
            margin: 1em 0;
            padding-left: 1em;
            color: #666;
            font-style: italic;
        }
        
        .mermaid {
            text-align: center;
            margin: 2em 0;
        }
        
        .document-footer {
            margin-top: 3em;
            padding-top: 20px;
            border-top: 1px solid #e2e8f0;
            text-align: center;
            color: #666;
            font-size: 0.9em;
        }
        
        @media (max-width: 768px) {
            .container {
                margin: 10px;
                padding: 15px;
            }
            
            .meta-info {
                flex-direction: column;
                gap: 10px;
            }
        }
        """
    
    def _parse_markdown_to_docx(self, doc: "Document", content: str):
        """Parse Markdown content and add it to the Word document"""
        lines = content.split('\n')
        
        for line in lines:
            line = line.strip()
            
            if not line:
                continue
                
            # Heading processing
            if line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                title_text = line.lstrip('#').strip()
                if level <= 6:
                    doc.add_heading(title_text, level)
                    continue
            
            # Code block processing (simplified)
            if line.startswith('```'):
                continue
                
            # List processing
            if line.startswith('- ') or line.startswith('* '):
                text = line[2:].strip()
                para = doc.add_paragraph(text, style='List Bullet')
                continue
                
            if re.match(r'^\d+\.', line):
                text = re.sub(r'^\d+\.\s*', '', line)
                para = doc.add_paragraph(text, style='List Number')
                continue
            
            # Normal paragraph
            if line:
                # Simple bold and italic processing
                line = re.sub(r'\*\*(.*?)\*\*', r'\1', line)  # Remove bold markers, can be manually set in Word later
                line = re.sub(r'\*(.*?)\*', r'\1', line)      # Remove italic markers
                doc.add_paragraph(line)
    
    def _export_pdf_reportlab(self, content: str, metadata: Optional[Dict] = None) -> bytes:
        """Export PDF using ReportLab"""
        try:
            buffer = io.BytesIO()
            
            # Create PDF document
            doc = SimpleDocTemplate(
                buffer,
                pagesize=A4,
                topMargin=1*inch,
                bottomMargin=1*inch,
                leftMargin=1*inch,
                rightMargin=1*inch
            )
            
            # Style settings
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Title'],
                fontSize=20,
                spaceAfter=30,
                alignment=1  # Centered
            )
            
            # Build content
            story = []
            
            # Add title
            title = metadata.get('title', 'VibeDoc Development Plan') if metadata else 'VibeDoc Development Plan'
            story.append(Paragraph(title, title_style))
            story.append(Spacer(1, 20))
            
            # Add metadata info
            meta_text = f"""
            Author: {metadata.get('author', 'VibeDoc AI Agent') if metadata else 'VibeDoc AI Agent'}<br/>
            Generated Time: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}<br/>
            Generated Tool: VibeDoc AI Agent
            """
            story.append(Paragraph(meta_text, styles['Normal']))
            story.append(Spacer(1, 30))
            
            # Simple processing of Markdown content (basic version)
            lines = content.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    story.append(Spacer(1, 12))
                    continue
                    
                if line.startswith('#'):
                    # Heading
                    level = len(line) - len(line.lstrip('#'))
                    title_text = line.lstrip('#').strip()
                    if level == 1:
                        story.append(Paragraph(title_text, styles['Heading1']))
                    elif level == 2:
                        story.append(Paragraph(title_text, styles['Heading2']))
                    else:
                        story.append(Paragraph(title_text, styles['Heading3']))
                else:
                    # Normal paragraph
                    story.append(Paragraph(line, styles['Normal']))
                    
                story.append(Spacer(1, 6))
            
            # Generate PDF
            doc.build(story)
            buffer.seek(0)
            
            logger.info("✅ PDF export successful (ReportLab)")
            return buffer.getvalue()
            
        except Exception as e:
            logger.error(f"❌ ReportLab PDF export failed: {e}")
            raise

# Global export manager instance
export_manager = ExportManager()